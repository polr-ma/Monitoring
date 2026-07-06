"""
ASR 审计日志模块 — 将所有语音识别结果写入 Word 文档
用于调试识别不准、漏识别等问题
"""

import os
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


@dataclass
class ASRAuditEntry:
    """单次 ASR 识别记录"""
    timestamp: datetime = field(default_factory=datetime.now)
    text: str = ""                          # 识别出的文本
    audio_duration_sec: float = 0.0         # 音频片段时长
    audio_peak: float = 0.0                 # 峰值电平
    buffer_chunks: int = 0                  # 缓冲段数
    matched_words: str = ""                 # 命中的违禁词（逗号分隔）
    anomaly_flags: str = ""                 # 异常标记（逗号分隔）


class ASRAuditLogger:
    """ASR 识别审计日志器 — 将所有识别结果写入 Word 文档"""

    _SAVE_INTERVAL = 10   # 每 3 条自动保存一次

    def __init__(self, output_dir: str = '.'):
        self._output_dir = output_dir
        self._doc: Optional[Document] = None
        self._table = None
        self._counter = 0
        self._pending = 0
        self._filepath: str = ''
        self._init_document()

    @property
    def filepath(self) -> str:
        return self._filepath

    # ── 文档初始化 ──────────────────────────────────────

    def _init_document(self):
        self._doc = Document()

        section = self._doc.sections[0]
        section.page_width = Cm(29.7)   # A4 横向
        section.page_height = Cm(21)

        title = self._doc.add_heading('ASR 语音识别审计日志', level=1)
        title.alignment = WD_TABLE_ALIGNMENT.CENTER

        now = datetime.now()
        self._doc.add_paragraph(
            f'生成时间：{now.strftime("%Y-%m-%d %H:%M:%S")}    '
            f'模型：SenseVoiceSmall'
        )
        self._doc.add_paragraph(
            '用途：记录每一次语音识别的完整结果，用于排查识别不准、漏识别等问题。'
        )
        self._doc.add_paragraph('')

        # 8 列
        headers = [
            '序号', '时间', '识别文本', '音频时长(s)',
            '峰值电平', '缓冲段数', '命中违禁词', '异常标记',
        ]
        self._table = self._doc.add_table(rows=1, cols=len(headers), style='Table Grid')
        self._table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = self._table.rows[0].cells
        for i, text in enumerate(headers):
            header_cells[i].text = text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 列宽（cm）
        widths = [
            Cm(1.0),   # 序号
            Cm(2.0),   # 时间
            Cm(9.0),   # 识别文本
            Cm(2.0),   # 音频时长
            Cm(2.0),   # 峰值电平
            Cm(1.8),   # 缓冲段数
            Cm(5.0),   # 命中违禁词
            Cm(4.0),   # 异常标记
        ]
        for i, width in enumerate(widths):
            for cell in self._table.columns[i].cells:
                cell.width = width

        filename = f'ASR审计_{now.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
        self._filepath = os.path.join(self._output_dir, filename)
        self._save()

    # ── 写入记录 ────────────────────────────────────────

    def add_entry(self, entry: ASRAuditEntry):
        """添加一条 ASR 识别记录"""
        self._counter += 1
        self._pending += 1

        row = self._table.add_row()
        cells = row.cells

        cells[0].text = str(self._counter)
        cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        cells[1].text = entry.timestamp.strftime('%H:%M:%S')

        cells[2].text = entry.text if entry.text else '(空)'
        if not entry.text:
            for run in cells[2].paragraphs[0].runs:
                run.font.color.rgb = None  # 用默认色，但标记为空

        cells[3].text = f'{entry.audio_duration_sec:.2f}'
        cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        cells[4].text = f'{entry.audio_peak:.0f}'
        cells[4].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        cells[5].text = str(entry.buffer_chunks)
        cells[5].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        cells[6].text = entry.matched_words if entry.matched_words else '-'

        cells[7].text = entry.anomaly_flags if entry.anomaly_flags else '-'

        # 统一字号
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

        # 如果命中违禁词，整行标红
        if entry.matched_words:
            for cell in cells:
                shading = cell._tc.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'FFD6D6',
                    qn('w:val'): 'clear',
                })
                shading.append(shd)

        # 如果有异常标记，标黄
        if entry.anomaly_flags and not entry.matched_words:
            for cell in cells:
                shading = cell._tc.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'FFF3CD',
                    qn('w:val'): 'clear',
                })
                shading.append(shd)

        if self._pending >= self._SAVE_INTERVAL:
            self._save()
            self._pending = 0

    def _save(self):
        if not self._doc or not self._filepath:
            return
        import tempfile, shutil
        max_retries = 3
        for attempt in range(max_retries):
            try:
                # ?????????????
                fd, tmp_path = tempfile.mkstemp(
                    suffix='.docx', dir=os.path.dirname(self._filepath))
                os.close(fd)
                self._doc.save(tmp_path)
                shutil.move(tmp_path, self._filepath)
                return
            except PermissionError:
                if attempt < max_retries - 1:
                    import time
                    time.sleep(0.5)
                else:
                    # ???????????????
                    base, ext = os.path.splitext(self._filepath)
                    fallback = f'{base}_{attempt}{ext}'
                    self._doc.save(fallback)
                    self._filepath = fallback

    def close(self):
        self._save()
