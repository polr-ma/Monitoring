"""违规记录模块 — 自动写入 Word 文档（含违规截图）"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class ViolationRecorder:
    """违规事件 Word 文档记录器"""

    TYPE_LABELS = {
        'leave_post': '离开工位',
        'return_post': '回到工位',
        'look_around': '东张西望',
        'head_down': '低头(瞌睡/玩手机)',
        'sleeping': '睡觉/闭眼',
        'forbidden_word': '语音违禁词',
    }

    def __init__(self, output_dir: str = '.'):
        self._output_dir = output_dir
        self._doc = None
        self._table = None
        self._counter = 0
        self._save_interval = 5
        self._pending_count = 0
        self._filepath = ''
        self._init_document()

    def _init_document(self):
        """创建 Word 文档和表格"""
        self._doc = Document()

        section = self._doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        title = self._doc.add_heading('直播间违规记录', level=1)
        title.alignment = WD_TABLE_ALIGNMENT.CENTER

        now = datetime.now()
        self._doc.add_paragraph(f'生成时间：{now.strftime("%Y-%m-%d %H:%M:%S")}')
        self._doc.add_paragraph('')

        # 5 列: 序号 | 时间 | 违规类型 | 具体描述 | 截图
        self._table = self._doc.add_table(rows=1, cols=5, style='Table Grid')
        self._table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = self._table.rows[0].cells
        headers = ['序号', '时间', '违规类型', '具体描述', '截图']
        for i, text in enumerate(headers):
            header_cells[i].text = text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

        widths = [Cm(1.2), Cm(2.8), Cm(2.8), Cm(7.0), Cm(5.2)]
        for i, width in enumerate(widths):
            for cell in self._table.columns[i].cells:
                cell.width = width

        filename = f'违规记录_{now.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
        self._filepath = os.path.join(self._output_dir, filename)
        self._save()

    def add_violation(self, event) -> str:
        """添加一条违规记录（含截图），返回文件路径"""
        self._counter += 1
        self._pending_count += 1

        row = self._table.add_row()
        cells = row.cells

        cells[0].text = str(self._counter)
        cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        cells[1].text = event.timestamp.strftime('%H:%M:%S')

        cells[2].text = self.TYPE_LABELS.get(event.violation_type, event.violation_type)

        cells[3].text = event.description

        for cell in cells[:4]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

        # ── 嵌入截图 ──
        if event.screenshot_path and os.path.exists(event.screenshot_path):
            try:
                para = cells[4].paragraphs[0]
                para.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = para.add_run()
                run.add_picture(event.screenshot_path, width=Cm(4.5))
            except Exception:
                cells[4].text = '(截图失败)'
        else:
            cells[4].text = '(无截图)'

        for paragraph in cells[4].paragraphs:
            for run in paragraph.runs:
                if not hasattr(run, '_r') or run._r.find(qn('wp:inline')) is None:
                    run.font.size = Pt(9)

        # 调整行高适配截图
        try:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = trPr.makeelement(qn('w:trHeight'), {
                qn('w:val'): '2000',
                qn('w:hRule'): 'atLeast',
            })
            trPr.insert(0, trHeight)
        except Exception:
            pass

        if self._pending_count >= self._save_interval:
            self._save()
            self._pending_count = 0

        return self._filepath

    def _save(self):
        """保存文档到磁盘"""
        if self._doc and self._filepath:
            self._doc.save(self._filepath)

    def close(self):
        """关闭时强制保存"""
        self._save()

    @property
    def filepath(self) -> str:
        return self._filepath
